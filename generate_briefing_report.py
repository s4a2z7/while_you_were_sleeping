"""
당신이 잠든 사이 브리핑 Word 문서 생성
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# 경로 설정
OUTPUT_DIR = "output/reports"
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "briefing_2025-12-05.docx")

# 문서 생성
doc = Document()

# 문서 여백 설정
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ========== 제목 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("당신이 잠든 사이")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(16, 185, 129)  # 초록색

# 날짜
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run("2025.12.05")
date_run.font.size = Pt(14)
date_run.font.color.rgb = RGBColor(156, 163, 175)

# 구분선
doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

# 빈 줄
doc.add_paragraph()

# ========== 오늘의 화제 종목 ==========
section1 = doc.add_heading("오늘의 화제 종목", level=1)
section1.runs[0].font.color.rgb = RGBColor(16, 185, 129)

# 종목명
ticker_para = doc.add_paragraph()
ticker_run = ticker_para.add_run("TESLA (TSLA)")
ticker_run.font.size = Pt(16)
ticker_run.font.bold = True

# 주가 정보 표
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'

# 헤더
header_cells = table.rows[0].cells
header_cells[0].text = "항목"
header_cells[1].text = "값"

# 데이터 행
data_rows = [
    ["주가", "$385.20"],
    ["변동률", "+8.7% (상승)"],
]

for i, (label, value) in enumerate(data_rows, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = label
    row_cells[1].text = value
    
    # 변동률 셀의 텍스트를 초록색으로
    if "상승" in value:
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(16, 185, 129)

# 선정 기준
criteria_para = doc.add_paragraph()
criteria_run = criteria_para.add_run("선정 기준")
criteria_run.font.bold = True
criteria_run.font.size = Pt(12)

criteria_list = doc.add_paragraph("거래량 1위", style='List Bullet')

doc.add_paragraph()

# ========== 왜 화제인가? ==========
section2 = doc.add_heading("왜 화제인가?", level=1)
section2.runs[0].font.color.rgb = RGBColor(16, 185, 129)

reasons = [
    "사이버트럭 판매량 급증",
    "FSD v13 출시 예고"
]

for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_paragraph()

# ========== 관련 뉴스 TOP 3 ==========
section3 = doc.add_heading("관련 뉴스 TOP 3", level=1)
section3.runs[0].font.color.rgb = RGBColor(16, 185, 129)

news_items = [
    ("테슬라 사이버트럭 미국 판매 3위", "Reuters"),
    ("FSD v13 무감독 자율주행", "Bloomberg"),
    ("중국 시장 반등", "CNBC"),
]

for i, (title, source) in enumerate(news_items, start=1):
    # 뉴스 제목
    news_para = doc.add_paragraph()
    num_run = news_para.add_run(f"{i}. ")
    num_run.font.bold = True
    num_run.font.color.rgb = RGBColor(16, 185, 129)
    
    title_run = news_para.add_run(title)
    title_run.font.bold = True
    
    # 출처
    source_para = doc.add_paragraph(f"출처: {source}", style='List Bullet 2')
    source_para.paragraph_format.left_indent = Inches(0.3)

doc.add_paragraph()

# ========== 하단 정보 ==========
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("While You Were Sleeping Dashboard")
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(107, 114, 128)

generated = doc.add_paragraph()
generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
generated_run = generated.add_run(f"생성 시간: {datetime.now().strftime('%Y.%m.%d %H:%M:%S')}")
generated_run.font.size = Pt(9)
generated_run.font.color.rgb = RGBColor(107, 114, 128)

# 문서 저장
doc.save(OUTPUT_PATH)
print(f"✅ Word 문서 생성 완료: {OUTPUT_PATH}")
print(f"📄 파일명: briefing_2025-12-05.docx")
print(f"📊 내용:")
print(f"  - 제목: 당신이 잠든 사이 | 2025.12.05")
print(f"  - 종목: TESLA (TSLA) | $385.20 (+8.7%)")
print(f"  - 섹션: 화제 종목, 왜 화제인가, 관련 뉴스 TOP 3")
